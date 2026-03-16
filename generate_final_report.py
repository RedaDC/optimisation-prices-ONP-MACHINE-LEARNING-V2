import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Configuration des chemins d'images
BRAIN_DIR = r"C:\Users\reda\.gemini\antigravity\brain\23d2232d-38de-4009-9f4c-4011492132ff"
OUTPUT_PATH = os.path.join(BRAIN_DIR, "RAPPORT_ANALYTIQUE_ONP_FINAL.docx")

IMAGES_CONFIG = [
    {
        "title": "CHAPITRE 1 : GESTION STRATÉGIQUE DES INFRASTRUCTURES",
        "sections": [
            {
                "subtitle": "1.2.1 Gestion des infrastructures",
                "image": "strategic_map_1772712844516.png",
                "legend": "Figure 1.1 : Visualisation cartographique du maillage territorial des ports de l'ONP (Interface Command Center)."
            },
            {
                "subtitle": "1.3 Problématique de la volatilité",
                "image": "market_pulse_1772712481065.png",
                "legend": "Figure 1.2 : Indicateurs de dynamisme des halles pour le suivi des flux de débarquement en temps réel."
            }
        ]
    },
    {
        "title": "CHAPITRE 2 : ANALYSE FINANCIÈRE ET DONNÉES",
        "description": "L'objectif est de montrer la rigueur de votre traitement de données et l'analyse comparative.",
        "sections": [
            {
                "subtitle": "2.2 Évolution du Chiffre d'Affaires",
                "image": "comparative_dashboard_1772713229713.png",
                "legend": "Figure 2.1 : Analyse comparative du Chiffre d'Affaires national et régional (Période 2024 vs 2025)."
            },
            {
                "subtitle": "2.3 Décomposition mathématique",
                "image": "price_volume_effect_1772714160978.png",
                "legend": "Figure 2.2 : Décomposition analytique de la variation du CA par les effets Prix et Volume."
            }
        ]
    },
    {
        "title": "CHAPITRE 3 : INTELLIGENCE DE MARCHÉ (ANALYTICS)",
        "description": "L'objectif est de montrer la compréhension statistique des prix.",
        "sections": [
            {
                "subtitle": "3.1 Volatilité",
                "image": "boxplots_volatility_final_1772714604132.png",
                "legend": "Figure 3.1 : Analyse de la volatilité et de la dispersion des cours par espèce halieutique."
            },
            {
                "subtitle": "3.3 Cycles de marché",
                "image": "seasonality_curves_final_1772714630497.png",
                "legend": "Figure 3.2 : Profils de saisonnalité historique des prix moyens de vente."
            }
        ]
    },
    {
        "title": "CHAPITRE 4 : MACHINE LEARNING",
        "description": "L'objectif est de montrer l'aspect prédictif et l'aide à la décision.",
        "sections": [
            {
                "subtitle": "4.1.2 Résultats du modèle",
                "image": "prediction_interface_sardine_1772715426411.png",
                "legend": "Figure 4.1 : Interface de prédiction assistée par XGBoost avec explication des facteurs d'influence."
            },
            {
                "subtitle": "4.2.2 Optimisation",
                "image": "landing_advisor_sardine_1772715471972.png",
                "legend": "Figure 4.2 : Module d'optimisation pour la maximisation de la recette nette des armateurs."
            }
        ]
    },
    {
        "title": "CHAPITRE 5 : SIMULATEUR ET REPORTING",
        "description": "L'objectif est de montrer l'aspect prospectif et institutionnel.",
        "sections": [
            {
                "subtitle": "5.1.1 Simulation de scénarios",
                "image": "what_if_simulator_sardine_1772715549798.png",
                "legend": "Figure 5.1 : Simulateur de scénarios prospectifs permettant de mesurer l'impact des chocs d'offre sur les prix."
            },
            {
                "subtitle": "5.2 Digitalisation du reporting",
                "image": "reporting_export_1772715614653.png",
                "legend": "Figure 5.2 : Module d'automatisation du reporting institutionnel au format Word."
            }
        ]
    }
]

def generate_report():
    doc = Document()
    
    # Header styling
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "RAPPORT D'INTELLIGENCE HALIEUTIQUE - ONP"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title Page
    title = doc.add_heading("RAPPORT ANALYTIQUE ET STRATÉGIQUE", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"\nDate : {datetime.now().strftime('%d/%m/%Y')}\n")
    doc.add_paragraph("Ce rapport présente les fonctionnalités avancées de la plateforme d'optimisation des prix et d'intelligence de données de l'Office National des Pêches (ONP).\n")
    
    doc.add_page_break()
    
    # Chapters
    for chapter in IMAGES_CONFIG:
        print(f"Traitement du chapitre : {chapter['title']}")
        doc.add_heading(chapter["title"], level=1)
        if "description" in chapter:
            desc_p = doc.add_paragraph(chapter["description"])
            desc_p.italic = True
        
        for sec in chapter["sections"]:
            print(f"  Insertion de la section : {sec['subtitle']}")
            doc.add_heading(sec["subtitle"], level=2)
            
            img_path = os.path.join(BRAIN_DIR, sec["image"])
            if os.path.exists(img_path):
                print(f"    Image trouvée : {sec['image']}")
                doc.add_picture(img_path, width=Inches(6))
                
                # Legend
                legend_p = doc.add_paragraph()
                legend_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = legend_p.add_run(sec["legend"])
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(80, 80, 80)
                run.italic = True
            else:
                print(f"    ERREUR : Image manquante : {img_path}")
                doc.add_paragraph(f"[ERREUR : Image manquante - {sec['image']}]")
            
            doc.add_paragraph() # Spacer
            
    # Save
    print(f"Sauvegarde du rapport vers : {OUTPUT_PATH}")
    doc.save(OUTPUT_PATH)
    print(f"Rapport généré avec succès.")

if __name__ == "__main__":
    generate_report()
