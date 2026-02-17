"""
Générateur de Rapport Word Professionnel pour Extraction 2024-2025
===================================================================
Crée un rapport complet avec logo ONP, tableaux, et graphiques embarqués
"""
import os
from datetime import datetime
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd

from extraction_2024_2025 import extract_summary_data, extract_vente_data, calculate_global_kpis

def create_extraction_word_report(output_path="Rapport_Extraction_2024_2025.docx"):
    """
    Génère un rapport Word professionnel complet
    
    Args:
        output_path: Chemin du fichier de sortie
        
    Returns:
        str: Chemin du fichier généré
    """
    doc = Document()
    
    # Style par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # ==================== HEADER AVEC LOGO ====================
    table_header = doc.add_table(rows=1, cols=2)
    table_header.width = Inches(6.5)
    cells = table_header.rows[0].cells
    
    logo_path = "logo onp.png"
    if os.path.exists(logo_path):
        run = cells[0].paragraphs[0].add_run()
        run.add_picture(logo_path, width=Inches(1.2))
    else:
        cells[0].text = "ONP"
    
    p_info = cells[1].paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_title = p_info.add_run("RAPPORT D'ANALYSE COMPARATIVE\n")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    
    p_info.add_run(f"EXTRACTION 2024-2025\n")
    p_info.add_run(f"DATE : {datetime.now().strftime('%d/%m/%Y')}\n")
    p_info.add_run(f"REF : ONP-EXT-{datetime.now().strftime('%Y%m%d')}")
    
    doc.add_paragraph("\n")
    
    # ==================== OBJET ====================
    p_obj = doc.add_paragraph()
    run_obj = p_obj.add_run("OBJET : Analyse Comparative des Performances 2024-2025 par DR et Espece")
    run_obj.bold = True
    run_obj.underline = True
    
    doc.add_paragraph("\n")
    
    # ==================== SECTION 1: SYNTHESE GLOBALE ====================
    doc.add_heading("1. SYNTHESE GLOBALE", level=2)
    
    kpis = calculate_global_kpis()
    
    table_kpi = doc.add_table(rows=1, cols=2)
    table_kpi.style = 'Table Grid'
    hdr = table_kpi.rows[0].cells
    hdr[0].text = "INDICATEUR"
    hdr[1].text = "VALEUR"
    for i in range(2):
        if hdr[i].paragraphs:
            hdr[i].paragraphs[0].runs[0].bold = True
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'EEEEEE')
            hdr[i]._tc.get_or_add_tcPr().append(shading_elm)
    
    kpi_items = [
        ("Chiffre d'Affaires Total 2024", f"{kpis['ca_2024_total_mdh']:,.2f} MDH"),
        ("Chiffre d'Affaires Total 2025", f"{kpis['ca_2025_total_mdh']:,.2f} MDH"),
        ("Variation du CA", f"{kpis['var_ca_mdh']:+,.2f} MDH ({kpis['var_ca_pct']:+.2f}%)"),
        ("Volume Total 2024", f"{kpis['volume_2024_total_t']:,.0f} Tonnes"),
        ("Volume Total 2025", f"{kpis['volume_2025_total_t']:,.0f} Tonnes"),
        ("Variation du Volume", f"{kpis['var_volume_pct']:+.2f}%"),
        ("Nombre de Delegations Regionales", f"{int(kpis['nb_dr'])}"),
        ("Nombre d'Especes", f"{int(kpis['nb_especes'])}"),
    ]
    
    for label, val in kpi_items:
        row = table_kpi.add_row().cells
        row[0].text = label
        row[1].text = val
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n")
    
    # ==================== GRAPHIQUE 1: Evolution CA ====================
    doc.add_paragraph("1.1 Evolution du Chiffre d'Affaires").bold = True
    
    try:
        fig_ca = go.Figure(data=[
            go.Bar(name='2024', x=['CA Total'], y=[kpis['ca_2024_total_mdh']], marker_color='#3B82F6'),
            go.Bar(name='2025', x=['CA Total'], y=[kpis['ca_2025_total_mdh']], marker_color='#10B981')
        ])
        fig_ca.update_layout(
            title='Comparaison CA 2024 vs 2025 (MDH)',
            yaxis_title='Chiffre d\'Affaires (MDH)',
            barmode='group',
            height=400
        )
        img_bytes = pio.to_image(fig_ca, format="png", width=1000, height=400, scale=2)
        doc.add_picture(BytesIO(img_bytes), width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"[Erreur export graphique: {str(e)}]")
    
    doc.add_paragraph("\n")
    
    # ==================== SECTION 2: ANALYSE PAR DR ====================
    doc.add_heading("2. ANALYSE PAR DELEGATION REGIONALE", level=2)
    
    df_summary = extract_summary_data()
    df_by_dr = df_summary.groupby('dr').agg({
        'ca_2024_kdh': 'sum',
        'ca_2025_kdh': 'sum',
        'var_ca_kdh': 'sum',
        'volume_2024_t': 'sum',
        'volume_2025_t': 'sum'
    }).reset_index()
    df_by_dr['var_ca_pct'] = (df_by_dr['var_ca_kdh'] / df_by_dr['ca_2024_kdh'] * 100).fillna(0)
    df_by_dr = df_by_dr.sort_values('var_ca_kdh', ascending=False)
    
    table_dr = doc.add_table(rows=1, cols=5)
    table_dr.style = 'Table Grid'
    headers_dr = ['DR', 'CA 2024 (MDH)', 'CA 2025 (MDH)', 'Variation (MDH)', 'Variation (%)']
    hdr_cells = table_dr.rows[0].cells
    for i, h in enumerate(headers_dr):
        hdr_cells[i].text = h
        if hdr_cells[i].paragraphs:
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'EEEEEE')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
    
    for _, row in df_by_dr.head(20).iterrows():
        row_cells = table_dr.add_row().cells
        row_cells[0].text = str(row['dr'])[:30]
        row_cells[1].text = f"{row['ca_2024_kdh']/1000:,.2f}"
        row_cells[2].text = f"{row['ca_2025_kdh']/1000:,.2f}"
        row_cells[3].text = f"{row['var_ca_kdh']/1000:+,.2f}"
        row_cells[4].text = f"{row['var_ca_pct']:+.2f}%"
    
    doc.add_paragraph("\n")
    
    # ==================== SECTION 3: TOP 20 ESPECES ====================
    doc.add_heading("3. TOP 20 ESPECES PAR CHIFFRE D'AFFAIRES 2025", level=2)
    
    df_by_espece = df_summary.groupby('espece').agg({
        'ca_2024_kdh': 'sum',
        'ca_2025_kdh': 'sum',
        'var_ca_kdh': 'sum',
        'volume_2024_t': 'sum',
        'volume_2025_t': 'sum'
    }).reset_index()
    df_by_espece['var_ca_pct'] = (df_by_espece['var_ca_kdh'] / df_by_espece['ca_2024_kdh'] * 100).fillna(0)
    df_by_espece = df_by_espece.sort_values('ca_2025_kdh', ascending=False)
    
    table_esp = doc.add_table(rows=1, cols=5)
    table_esp.style = 'Table Grid'
    headers_esp = ['Espece', 'CA 2024 (MDH)', 'CA 2025 (MDH)', 'Variation (MDH)', 'Variation (%)']
    hdr_cells_esp = table_esp.rows[0].cells
    for i, h in enumerate(headers_esp):
        hdr_cells_esp[i].text = h
        if hdr_cells_esp[i].paragraphs:
            hdr_cells_esp[i].paragraphs[0].runs[0].bold = True
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'EEEEEE')
            hdr_cells_esp[i]._tc.get_or_add_tcPr().append(shading_elm)
    
    for _, row in df_by_espece.head(20).iterrows():
        row_cells = table_esp.add_row().cells
        row_cells[0].text = str(row['espece'])[:40]
        row_cells[1].text = f"{row['ca_2024_kdh']/1000:,.2f}"
        row_cells[2].text = f"{row['ca_2025_kdh']/1000:,.2f}"
        row_cells[3].text = f"{row['var_ca_kdh']/1000:+,.2f}"
        row_cells[4].text = f"{row['var_ca_pct']:+.2f}%"
    
    doc.add_paragraph("\n")
    
    # ==================== GRAPHIQUE 2: Top 10 Especes ====================
    doc.add_paragraph("3.1 Top 10 Especes - Variation du CA").bold = True
    
    try:
        top_10_esp = df_by_espece.head(10)
        fig_esp = px.bar(
            top_10_esp,
            x='espece',
            y='var_ca_kdh',
            title='Top 10 Especes - Variation du CA (KDh)',
            labels={'var_ca_kdh': 'Variation CA (KDh)', 'espece': 'Espece'},
            color='var_ca_kdh',
            color_continuous_scale='RdYlGn'
        )
        fig_esp.update_layout(height=500)
        img_bytes_esp = pio.to_image(fig_esp, format="png", width=1000, height=500, scale=2)
        doc.add_picture(BytesIO(img_bytes_esp), width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"[Erreur export graphique: {str(e)}]")
    
    doc.add_paragraph("\n")
    
    # ==================== FOOTER ====================
    doc.add_paragraph("\n" + "-" * 80)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run("Office National des Peches - Departement de l'Intelligence de Donnees\n")
    run_footer.bold = True
    p_footer.add_run("www.onp.ma | Support Technique : master.data@onp.ma")
    
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    print("="*80)
    print("GENERATION DU RAPPORT WORD")
    print("="*80)
    
    output_file = create_extraction_word_report()
    print(f"\n[OK] Rapport genere avec succes: {output_file}")
