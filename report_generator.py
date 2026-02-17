"""
Générateur de Rapport Word Professionnel (Style Facture/Institutionnel)
=====================================================================

Crée un document .docx avec une mise en page formelle incluant:
- Logo ONP
- Numérotation officielle
- Tableaux de métriques structurés
- Synthèse stratégique
- Graphiques Plotly exportés
"""

import os
from datetime import datetime
import plotly.io as pio
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_institutional_word_report(df_metrics, filters=None, df_detailed=None, output_path="Rapport_ONP_Excellence.docx"):
    """
    Crée un rapport institutionnel standard.
    """
    doc = Document()
    
    # --- Configuration des Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # --- HEADER SECTION (LOGO + INFO) ---
    table_header = doc.add_table(rows=1, cols=2)
    table_header.width = Inches(6.5)
    
    # Logo (si présent)
    cells = table_header.rows[0].cells
    logo_path = "logo onp.png"
    if os.path.exists(logo_path):
        run = cells[0].paragraphs[0].add_run()
        run.add_picture(logo_path, width=Inches(1.2))
    
    # Numéro et Date (Style Facture)
    p_info = cells[1].paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_title = p_info.add_run("RAPPORT D'INTELLIGENCE\n")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(15, 23, 42) # Navy
    
    p_info.add_run(f"N° RÉF : ONP-{datetime.now().strftime('%Y%m%d')}-A1\n")
    p_info.add_run(f"DATE : {datetime.now().strftime('%d/%m/%Y')}\n")
    p_info.add_run("ÉMIS PAR : Plateforme Stratégique Halieutis")
    
    doc.add_paragraph("\n")
    
    # --- OBJET DU RAPPORT ---
    p_subject = doc.add_paragraph()
    run_subject = p_subject.add_run("OBJET : SYNTHÈSE ANALYTIQUE ET PRÉVISIONS DU MARCHÉ HALIEUTIQUE")
    run_subject.bold = True
    run_subject.underline = True
    
    # --- PÉRIMÈTRE DE L'ANALYSE ---
    if filters:
        doc.add_paragraph("\nPORTÉE DE L'ANALYSE :")
        p_scope = doc.add_paragraph()
        p_scope.add_run("Ports sélectionnés : ").bold = True
        p_scope.add_run(", ".join(filters.get('ports', [])) if filters.get('ports') else "Tous les ports")
        
        p_scope2 = doc.add_paragraph()
        p_scope2.add_run("Espèces sélectionnées : ").bold = True
        p_scope2.add_run(", ".join(filters.get('species', [])) if filters.get('species') else "Toutes les espèces")
    
    doc.add_paragraph("\n")
    
    # --- TABLE DES MÉTRIQUES (STYLE ITEMS FACTURE) ---
    doc.add_paragraph("INDICATEURS DE PERFORMANCE DU PÉRIMÈTRE :")
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    # En-têtes avec fond légèrement coloré (gris)
    headers = ["DÉSIGNATION DE L'ANALYSE", "VALEUR", "ÉTAT / REMARQUE"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'EEEEEE')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    # Données des métriques
    items = [
        ("Recette Totale du Périmètre (MDH)", f"{df_metrics.get('recette_totale_mdh', 0):.2f} MDH", "Validé"),
        ("Volume Total Débarqué (Tonnes)", f"{df_metrics.get('volume_total_tonnes', 0):,.0f} T", "Conforme"),
        ("Prix Moyen Pondéré (DH/kg)", f"{df_metrics.get('prix_moyen_dh_kg', 0):.2f} DH/kg", "Observation"),
        ("Espèce Majoritaire", f"{df_metrics.get('espece_plus_rentable', 'N/A')}", "Segment Clé"),
    ]
    
    for desc, val, rem in items:
        row_cells = table.add_row().cells
        row_cells[0].text = desc
        row_cells[1].text = val
        row_cells[2].text = rem
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- SYNTHÈSE STRATÉGIQUE ---
    doc.add_heading("SYNTHÈSE EXÉCUTIVE", level=2)
    p_syn = doc.add_paragraph()
    p_syn.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    scope_desc = "sur l'ensemble des ports" if not (filters and filters.get('ports')) else f"sur les ports de {', '.join(filters['ports'])}"
    p_syn.add_run(
        f"L'analyse approfondie effectuée {scope_desc} démontre une dynamique de marché robuste. "
        "L'optimisation des prix via nos modèles prédictifs permet de sécuriser les revenus des pêcheurs "
        "tout en assurant une disponibilité constante des ressources halieutiques. "
    )
    
    doc.add_paragraph("\n")
    
    # --- FOOTER / SIGNATURE ---
    doc.add_paragraph("-" * 80)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run("Office National des Pêches - Département de l'Intelligence de Données\n")
    run_footer.bold = True
    p_footer.add_run("www.onp.ma | Support Technique : master.data@onp.ma")
    
    # Save to buffer if no output_path provided (for testing/integration)
    doc.save(output_path)
    return output_path

def create_reduction_word_report(df_reduction, stats, plotly_figs=None):
    """
    Génère un rapport Word pour l'analyse de diminution du CA 2024-2025.
    df_reduction: DataFrame avec les données traitées
    stats: dictionnaire des KPIs globaux
    plotly_figs: dictionnaire de figures plotly {title: fig}
    """
    from io import BytesIO
    doc = Document()
    
    # Style par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # Header Section
    table_header = doc.add_table(rows=1, cols=2)
    table_header.width = Inches(6.5)
    cells = table_header.rows[0].cells
    
    logo_path = "logo onp.png"
    if os.path.exists(logo_path):
        run = cells[0].paragraphs[0].add_run()
        run.add_picture(logo_path, width=Inches(1.2))
    
    p_info = cells[1].paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_title = p_info.add_run("RAPPORT D'ANALYSE COMPARATIVE\n")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    
    p_info.add_run(f"PÉRIODE : 2024-2025\n")
    p_info.add_run(f"DATE : {datetime.now().strftime('%d/%m/%Y')}\n")
    p_info.add_run("OFFICE NATIONAL DES PÊCHES")
    
    doc.add_paragraph("\n")
    
    # Objet
    p_obj = doc.add_paragraph()
    run_obj = p_obj.add_run("OBJET : Analyse de l'évolution du Chiffre d'Affaires par Port, Délégation et Espèce")
    run_obj.bold = True
    run_obj.underline = True
    
    doc.add_paragraph("\n")
    
    # Résumé Global (KPIs)
    doc.add_heading("1. RÉSUMÉ GLOBAL DES INDICATEURS", level=2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "INDICATEUR"
    hdr[1].text = "VALEUR"
    for i in range(2): 
        if hdr[i].paragraphs:
            hdr[i].paragraphs[0].runs[0].bold = True

    kpis = [
        ("Chiffre d'Affaires Total 2024", f"{stats['ca_2024']/1000:,.2f} MDH"),
        ("Chiffre d'Affaires Total 2025", f"{stats['ca_2025']/1000:,.2f} MDH"),
        ("Volume Total Débarqué 2024", f"{stats.get('vol_2024', 0):,.2f} Tonnes"),
        ("Volume Total Débarqué 2025", f"{stats.get('vol_2025', 0):,.2f} Tonnes"),
        ("Variation CA Relative", f"{stats['diff_pct']:+.2f} %")
    ]
    
    for label, val in kpis:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = val

    doc.add_paragraph("\n")
    
    # Insertion des graphiques si disponibles
    if plotly_figs:
        doc.add_heading("2. ANALYSE VISUELLE (GRAPHIQUES)", level=2)
        for title, fig in plotly_figs.items():
            doc.add_paragraph(title).bold = True
            try:
                # Export plotly to image buffer
                img_bytes = pio.to_image(fig, format="png", width=1000, height=600, scale=2)
                doc.add_picture(BytesIO(img_bytes), width=Inches(6))
                doc.add_paragraph("\n")
            except Exception as e:
                doc.add_paragraph(f"[Erreur export graphique: {str(e)}]")
                print(f"Error exporting chart {title}: {e}")

    # Analyse par Port
    doc.add_heading("3. ANALYSE DÉTAILLÉE PAR PORT", level=2)
    df_port = df_reduction.groupby('port')[['ca_2024_kdh', 'ca_2025_kdh', 'ca_diff_kdh']].sum().sort_values('ca_diff_kdh')
    
    table_port = doc.add_table(rows=1, cols=4)
    table_port.style = 'Table Grid'
    hdrs = ["PORT", "CA 2024 (KDh)", "CA 2025 (KDh)", "VARIATION (KDh)"]
    for i, h in enumerate(hdrs):
        table_port.rows[0].cells[i].text = h
        if table_port.rows[0].cells[i].paragraphs:
            table_port.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
    for port, row in df_port.iterrows():
        cells = table_port.add_row().cells
        cells[0].text = str(port)
        cells[1].text = f"{row['ca_2024_kdh']:,.0f}".replace(',', ' ')
        cells[2].text = f"{row['ca_2025_kdh']:,.0f}".replace(',', ' ')
        cells[3].text = f"{row['ca_diff_kdh']:,.0f}".replace(',', ' ')

    doc.add_paragraph("\n")
    
    # Analyse par Délégation
    doc.add_heading("4. VARIATION PAR DÉLÉGATION", level=2)
    df_del = df_reduction.groupby('delegation')[['ca_2024_kdh', 'ca_2025_kdh', 'ca_diff_kdh']].sum().sort_values('ca_diff_kdh')
    
    table_del = doc.add_table(rows=1, cols=4)
    table_del.style = 'Table Grid'
    hdr_del = ["DÉLÉGATION", "CA 2024", "CA 2025", "VARIATION"]
    for i, h in enumerate(hdr_del):
        table_del.rows[0].cells[i].text = h
        if table_del.rows[0].cells[i].paragraphs:
            table_del.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
    for dele, row in df_del.iterrows():
        cells = table_del.add_row().cells
        cells[0].text = str(dele)
        cells[1].text = f"{row['ca_2024_kdh']:,.0f}".replace(',', ' ')
        cells[2].text = f"{row['ca_2025_kdh']:,.0f}".replace(',', ' ')
        cells[3].text = f"{row['ca_diff_kdh']:,.0f}".replace(',', ' ')

    doc.add_paragraph("\n")
    
    # Analyse par Espèce
    doc.add_heading("5. ANALYSE PAR ESPÈCE", level=2)
    df_esp = df_reduction.groupby('espece_categorie')[['ca_2024_kdh', 'ca_2025_kdh', 'ca_diff_kdh']].sum().sort_values('ca_diff_kdh')
    
    table_esp = doc.add_table(rows=1, cols=4)
    table_esp.style = 'Table Grid'
    hdr_esp = ["ESPÈCE", "CA 2024", "CA 2025", "VARIATION"]
    for i, h in enumerate(hdr_esp):
        table_esp.rows[0].cells[i].text = h
        if table_esp.rows[0].cells[i].paragraphs:
            table_esp.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
    for esp, row in df_esp.iterrows():
        cells = table_esp.add_row().cells
        cells[0].text = str(esp)[:40]
        cells[1].text = f"{row['ca_2024_kdh']:,.0f}".replace(',', ' ')
        cells[2].text = f"{row['ca_2025_kdh']:,.0f}".replace(',', ' ')
        cells[3].text = f"{row['ca_diff_kdh']:,.0f}".replace(',', ' ')

    # Signature
    doc.add_paragraph("\n" + "-" * 80)
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig.add_run("Office National des Pêches - Département de l'Intelligence de Données").bold = True
    
    # Save to memory buffer
    target_stream = BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()


def create_comparison_word_report(df_effects, output_path="Rapport_Comparaison_2024_2025.docx"):
    """
    Génère un rapport Word professionnel pour la comparaison 2024-2025 avec logo et graphiques.
    """
    import plotly.express as px
    import plotly.graph_objects as go
    
    doc = Document()
    
    # Style par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # --- HEADER AVEC LOGO ---
    table_header = doc.add_table(rows=1, cols=2)
    table_header.width = Inches(6.5)
    cells = table_header.rows[0].cells
    
    logo_path = "logo onp.png"
    if os.path.exists(logo_path):
        run = cells[0].paragraphs[0].add_run()
        run.add_picture(logo_path, width=Inches(1.2))
    else:
        cells[0].text = "ONP"
    
    p_info = cells[1].paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_title = p_info.add_run("RAPPORT D'ANALYSE COMPARATIVE\n")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    
    p_info.add_run(f"PERIODE : 2024-2025\n")
    p_info.add_run(f"DATE : {datetime.now().strftime('%d/%m/%Y')}\n")
    p_info.add_run(f"REF : ONP-DR-{datetime.now().strftime('%Y%m%d')}")
    
    doc.add_paragraph("\n")
    
    # --- OBJET ---
    p_obj = doc.add_paragraph()
    run_obj = p_obj.add_run("OBJET : Analyse des Effets Prix-Volume par Espece (2024 vs 2025)")
    run_obj.bold = True
    run_obj.underline = True
    
    doc.add_paragraph("\n")
    
    if df_effects.empty:
        doc.add_paragraph("Aucune donnee disponible pour la comparaison.")
        doc.save(output_path)
        return output_path
        
    # --- INDICATEURS CLES ---
    doc.add_heading("1. SYNTHESE GLOBALE", level=2)
    
    total_var = df_effects['variation_mdh'].sum()
    total_vol_eff = df_effects['effet_volume_mdh'].sum()
    total_pri_eff = df_effects['effet_prix_mdh'].sum()
    total_ca_2024 = df_effects['recette_2024_mdh'].sum()
    total_ca_2025 = df_effects['recette_2025_mdh'].sum()
    
    table_kpi = doc.add_table(rows=1, cols=2)
    table_kpi.style = 'Table Grid'
    hdr = table_kpi.rows[0].cells
    hdr[0].text = "INDICATEUR"
    hdr[1].text = "VALEUR"
    for i in range(2):
        if hdr[i].paragraphs:
            hdr[i].paragraphs[0].runs[0].bold = True
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'EEEEEE')
            hdr[i]._tc.get_or_add_tcPr().append(shading_elm)
    
    kpis = [
        ("Chiffre d'Affaires Total 2024", f"{total_ca_2024:,.2f} MDH"),
        ("Chiffre d'Affaires Total 2025", f"{total_ca_2025:,.2f} MDH"),
        ("Variation Totale du CA", f"{total_var:+,.2f} MDH"),
        ("Effet Volume Total", f"{total_vol_eff:+,.2f} MDH"),
        ("Effet Prix Total", f"{total_pri_eff:+,.2f} MDH"),
    ]
    
    for label, val in kpis:
        row = table_kpi.add_row().cells
        row[0].text = label
        row[1].text = val
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\\n")
    
    # --- METHODOLOGIE ---
    doc.add_heading("2. METHODOLOGIE DE CALCUL", level=2)
    doc.add_paragraph("L'analyse decompose la variation du CA en deux composantes :")
    doc.add_paragraph("- Effet Volume = (Volume_2025 - Volume_2024) * Prix_2024", style='List Bullet')
    doc.add_paragraph("- Effet Prix = (Prix_2025 - Prix_2024) * Volume_2025", style='List Bullet')
    doc.add_paragraph("- Variation Totale = Effet Volume + Effet Prix", style='List Bullet')
    
    doc.add_paragraph("\n")
    
    # --- GRAPHIQUES ---
    doc.add_heading("3. ANALYSE VISUELLE", level=2)
    
    # Graphique 1: Top 10 Especes par Variation
    doc.add_paragraph("3.1 Top 10 Especes par Variation de CA").bold = True
    try:
        top_10 = df_effects.nlargest(10, 'variation_mdh')
        fig1 = px.bar(
            top_10,
            x='espece',
            y='variation_mdh',
            title='Top 10 Especes - Variation du CA (MDH)',
            labels={'variation_mdh': 'Variation (MDH)', 'espece': 'Espece'},
            color='variation_mdh',
            color_continuous_scale='RdYlGn'
        )
        img_bytes1 = pio.to_image(fig1, format="png", width=1000, height=600, scale=2)
        doc.add_picture(BytesIO(img_bytes1), width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"[Erreur export graphique 1: {str(e)}]")
    
    doc.add_paragraph("\\n")
    
    # Graphique 2: Decomposition Effets Prix-Volume (Top 10)
    doc.add_paragraph("3.2 Decomposition Effets Prix-Volume (Top 10)").bold = True
    try:
        top_10_abs = df_effects.iloc[:10]
        df_plot = top_10_abs.melt(
            id_vars=['espece'],
            value_vars=['effet_volume_mdh', 'effet_prix_mdh'],
            var_name='Type Effet',
            value_name='Valeur (MDH)'
        )
        df_plot['Type Effet'] = df_plot['Type Effet'].replace({
            'effet_volume_mdh': 'Effet Volume',
            'effet_prix_mdh': 'Effet Prix'
        })
        
        fig2 = px.bar(
            df_plot,
            x='espece',
            y='Valeur (MDH)',
            color='Type Effet',
            barmode='group',
            title='Decomposition Prix-Volume (Top 10 Especes)',
            color_discrete_map={'Effet Volume': '#0EA5E9', 'Effet Prix': '#10B981'}
        )
        img_bytes2 = pio.to_image(fig2, format="png", width=1000, height=600, scale=2)
        doc.add_picture(BytesIO(img_bytes2), width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"[Erreur export graphique 2: {str(e)}]")
    
    doc.add_paragraph("\\n")
    
    # --- TABLEAU DETAILLE ---
    doc.add_heading("4. DETAIL PAR ESPECE (Top 30)", level=2)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    headers = ['Espece', 'CA 2024 (MDH)', 'CA 2025 (MDH)', 'Variation (MDH)', 'Effet Vol.', 'Effet Prix']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        if hdr_cells[i].paragraphs:
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'EEEEEE')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        
    # Remplissage avec les 30 premieres lignes
    for _, row in df_effects.head(30).iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['espece'])[:40]
        row_cells[1].text = f"{row['recette_2024_mdh']:,.2f}"
        row_cells[2].text = f"{row['recette_2025_mdh']:,.2f}"
        row_cells[3].text = f"{row['variation_mdh']:+,.2f}"
        row_cells[4].text = f"{row['effet_volume_mdh']:+,.2f}"
        row_cells[5].text = f"{row['effet_prix_mdh']:+,.2f}"
    
    # --- FOOTER ---
    doc.add_paragraph("\n" + "-" * 80)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run("Office National des Peches - Departement de l'Intelligence de Donnees\n")
    run_footer.bold = True
    p_footer.add_run("www.onp.ma | Support Technique : master.data@onp.ma")
    
    doc.save(output_path)
    return output_path
